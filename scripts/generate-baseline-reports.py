#!/usr/bin/env python3
"""
Generate baseline report bundle for project-state v2.0.
Outputs .docx and .xlsx files matching the ai26.10 quality standard.
"""

import os, sys, json, yaml, datetime, argparse, shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── CLI args ──
parser = argparse.ArgumentParser(description="Generate baseline report bundle from .project-state/")
parser.add_argument("--date", default=datetime.date.today().isoformat(), help="Report date (YYYY-MM-DD)")
parser.add_argument("--output-dir", default=None, help="Override output directory")
parser.add_argument("--state-dir", default=None, help="Path to .project-state/ (auto-detected if omitted)")
args = parser.parse_args()

TODAY = args.date

# ── Find .project-state/ ──
def find_state_dir():
    if args.state_dir:
        return Path(args.state_dir)
    # Walk up from script location
    d = Path(__file__).resolve().parent.parent
    for candidate in [d / ".project-state", d]:
        if (candidate / "manifest.yaml").exists():
            return candidate
    # Walk up from cwd
    d = Path.cwd()
    while d != d.parent:
        if (d / ".project-state" / "manifest.yaml").exists():
            return d / ".project-state"
        d = d.parent
    print("ERROR: No .project-state/manifest.yaml found", file=sys.stderr)
    sys.exit(1)

STATE = find_state_dir()

# ── Output directory ──
if args.output_dir:
    OUT = Path(args.output_dir)
else:
    OUT = STATE / "reports" / "baseline" / f"Baseline-Reports-{TODAY}"
OUT.mkdir(parents=True, exist_ok=True)

# ── Colors ──
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2E, 0x86, 0xAB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)

XL_NAVY = "1B2A4A"
XL_ACCENT = "2E86AB"
XL_LIGHT = "F0F4F8"
XL_WHITE = "FFFFFF"
XL_GREEN = "27AE60"
XL_AMBER = "F39C12"
XL_RED = "E74C3C"
XL_GRAY = "BDC3C7"

# ── Load state ──
def load_yaml(p):
    with open(p) as f:
        return yaml.safe_load(f)

def load_json(p):
    with open(p) as f:
        return json.load(f)

manifest = load_yaml(STATE / "manifest.yaml")
state = load_json(STATE / "state.json")

milestones = []
for p in sorted((STATE / "milestones").glob("*.yaml")):
    milestones.append(load_yaml(p))

risks = []
for p in sorted((STATE / "risks").glob("*.yaml")):
    risks.append(load_yaml(p))

# ── Helper: week number from date ──
def date_to_week(date_str, project_start="2026-04-24"):
    if not date_str:
        return None
    d = datetime.date.fromisoformat(str(date_str))
    s = datetime.date.fromisoformat(project_start)
    return max(1, ((d - s).days // 7) + 1)

# ── Helper: docx formatting ──
def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F0F4F8")
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): 'BDC3C7',
        })
        borders.append(el)
    tblPr.append(borders)
    return table

def doc_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def doc_para(doc, text, bold=False, style='Body Text'):
    p = doc.add_paragraph(text, style=style)
    if bold:
        for run in p.runs:
            run.font.bold = True
    return p

# ═══════════════════════════════════════════════════════════════
# 1. TRACKER XLSX
# ═══════════════════════════════════════════════════════════════

def build_tracker():
    wb = Workbook()

    header_font = Font(name="Calibri", bold=True, color=XL_WHITE, size=10)
    header_fill = PatternFill(start_color=XL_NAVY, end_color=XL_NAVY, fill_type="solid")
    accent_fill = PatternFill(start_color=XL_ACCENT, end_color=XL_ACCENT, fill_type="solid")
    light_fill = PatternFill(start_color=XL_LIGHT, end_color=XL_LIGHT, fill_type="solid")
    green_fill = PatternFill(start_color=XL_GREEN, end_color=XL_GREEN, fill_type="solid")
    amber_fill = PatternFill(start_color=XL_AMBER, end_color=XL_AMBER, fill_type="solid")
    red_fill = PatternFill(start_color=XL_RED, end_color=XL_RED, fill_type="solid")
    gray_fill = PatternFill(start_color=XL_GRAY, end_color=XL_GRAY, fill_type="solid")

    thin_border = Border(
        left=Side(style='thin', color='BDC3C7'),
        right=Side(style='thin', color='BDC3C7'),
        top=Side(style='thin', color='BDC3C7'),
        bottom=Side(style='thin', color='BDC3C7'),
    )
    wrap = Alignment(wrap_text=True, vertical='top')

    def style_header_row(ws, row, max_col):
        for c in range(1, max_col + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(wrap_text=True, vertical='center')
            cell.border = thin_border

    def style_data_cell(ws, row, col, alt=False):
        cell = ws.cell(row=row, column=col)
        cell.font = Font(name="Calibri", size=10)
        cell.alignment = wrap
        cell.border = thin_border
        if alt:
            cell.fill = light_fill
        return cell

    # ── Home sheet ──
    ws = wb.active
    ws.title = "Home"
    ws.sheet_properties.tabColor = XL_NAVY
    ws['A1'] = "project-state — Project Tracker v1.0"
    ws['A1'].font = Font(name="Calibri", bold=True, size=16, color=XL_NAVY)
    ws['A2'] = "Consolidated source-of-truth workbook · Milestones · Deliverables · Risks · Gantt"
    ws['A2'].font = Font(name="Calibri", size=11, color=XL_ACCENT)
    ws['A3'] = f"Issued {TODAY} · Generated from .project-state/ substrate"
    ws['A3'].font = Font(name="Calibri", italic=True, size=10)
    ws['A5'] = "Source sheets (edit these)"
    ws['A5'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)
    ws['A6'] = "Milestones — 10 milestones with status, %, dates, deliverables, owner"
    ws['A7'] = "Deliverables — 32 deliverables mapped to parent milestones"
    ws['A8'] = "Risks — 11 open risks with likelihood, impact, score, mitigation"
    ws['A9'] = "Phases — 4 agile-default lifecycle phases"
    ws['A11'] = "Derived views (read-only)"
    ws['A11'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)
    ws['A12'] = "Dashboard — live status summary computed from source sheets"
    ws['A13'] = "Monthly Gantt — 12-month Gantt derived from milestone start/end weeks"
    ws['A14'] = "Dependencies — milestone dependency map"
    ws['A15'] = "Legend — color codes, status values, conventions"
    ws.column_dimensions['A'].width = 80

    # ── Dashboard sheet ──
    ws = wb.create_sheet("Dashboard")
    ws.sheet_properties.tabColor = XL_ACCENT
    ws['A1'] = "project-state — Dashboard"
    ws['A1'].font = Font(name="Calibri", bold=True, size=14, color=XL_NAVY)
    ws['A2'] = "Live status computed from source sheets. Do not edit directly."
    ws['A2'].font = Font(name="Calibri", italic=True, size=10)

    ws['A4'] = "OVERALL PROJECT STATUS"
    ws['A4'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    labels = [
        ("Project", manifest['project']['name']),
        ("Phase", state['current_phase']),
        ("Pack", manifest['project']['packs_loaded'][0]),
        ("Start date", manifest['project']['start_date']),
        ("Lead", "David Olsson"),
        ("Generated", TODAY),
    ]
    for i, (k, v) in enumerate(labels):
        ws.cell(row=5+i, column=1, value=k).font = Font(name="Calibri", bold=True, size=10)
        ws.cell(row=5+i, column=2, value=v).font = Font(name="Calibri", size=10)

    ws['A13'] = "MILESTONE SUMMARY"
    ws['A13'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    complete = sum(1 for m in milestones if m.get('status') == 'complete')
    in_progress = sum(1 for m in milestones if m.get('status') == 'in_progress')
    planned = sum(1 for m in milestones if m.get('status') == 'planned')

    summary_labels = [
        ("Total milestones", len(milestones)),
        ("Complete", complete),
        ("In progress", in_progress),
        ("Planned", planned),
        ("Total deliverables", sum(len(m.get('deliverables', [])) for m in milestones)),
        ("Deliverables complete", sum(1 for m in milestones for d in m.get('deliverables', []) if d.get('status') == 'delivered')),
    ]
    for i, (k, v) in enumerate(summary_labels):
        ws.cell(row=14+i, column=1, value=k).font = Font(name="Calibri", bold=True, size=10)
        ws.cell(row=14+i, column=2, value=v).font = Font(name="Calibri", size=10)

    ws['A22'] = "RISK SUMMARY"
    ws['A22'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    critical = sum(1 for r in risks if r.get('likelihood') == 'high' and r.get('impact') == 'high')
    risk_labels = [
        ("Total open risks", len(risks)),
        ("Critical (high/high)", critical),
        ("High impact", sum(1 for r in risks if r.get('impact') == 'high')),
        ("Medium impact", sum(1 for r in risks if r.get('impact') == 'medium')),
    ]
    for i, (k, v) in enumerate(risk_labels):
        ws.cell(row=23+i, column=1, value=k).font = Font(name="Calibri", bold=True, size=10)
        ws.cell(row=23+i, column=2, value=v).font = Font(name="Calibri", size=10)

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 50

    # ── Milestones sheet ──
    ws = wb.create_sheet("Milestones")
    ws.sheet_properties.tabColor = XL_GREEN
    ws['A1'] = "Milestones — source of truth (edit here)"
    ws['A1'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    headers = ["ID", "Owner", "Title", "Status", "% Complete", "Planned Start",
               "Planned End", "Actual Start", "Actual End", "Start Week", "End Week",
               "Technical Progress", "Deliverable Count", "Deliverables Complete"]
    row = 3
    for ci, h in enumerate(headers, 1):
        ws.cell(row=row, column=ci, value=h)
    style_header_row(ws, row, len(headers))

    for ri, m in enumerate(milestones):
        r = row + 1 + ri
        alt = ri % 2 == 1
        vals = [
            m['id'],
            m.get('owner_person', ''),
            m.get('title', ''),
            m.get('status', ''),
            m.get('percent_complete', 0),
            str(m.get('planned_start', '')),
            str(m.get('planned_end', '')),
            str(m.get('actual_start', '') or ''),
            str(m.get('actual_end', '') or ''),
            date_to_week(m.get('planned_start')),
            date_to_week(m.get('planned_end')),
            m.get('technical_progress', ''),
            len(m.get('deliverables', [])),
            sum(1 for d in m.get('deliverables', []) if d.get('status') == 'delivered'),
        ]
        for ci, v in enumerate(vals, 1):
            cell = style_data_cell(ws, r, ci, alt)
            cell.value = v
            # Color status
            if ci == 4:
                if v == 'complete':
                    cell.fill = green_fill
                    cell.font = Font(name="Calibri", size=10, color=XL_WHITE, bold=True)
                elif v == 'in_progress':
                    cell.fill = amber_fill
                    cell.font = Font(name="Calibri", size=10, color=XL_WHITE, bold=True)

    widths = [22, 22, 50, 12, 12, 14, 14, 14, 14, 10, 10, 50, 14, 16]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── Deliverables sheet ──
    ws = wb.create_sheet("Deliverables")
    ws.sheet_properties.tabColor = XL_GREEN
    ws['A1'] = "Deliverables — source of truth (edit here)"
    ws['A1'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    headers = ["Deliverable ID", "Milestone", "Title", "Status", "Milestone Status", "Milestone % Complete"]
    row = 3
    for ci, h in enumerate(headers, 1):
        ws.cell(row=row, column=ci, value=h)
    style_header_row(ws, row, len(headers))

    dr = row + 1
    for m in milestones:
        for d in m.get('deliverables', []):
            alt = (dr - row - 1) % 2 == 1
            vals = [d['id'], m['id'], d['title'], d['status'], m['status'], m.get('percent_complete', 0)]
            for ci, v in enumerate(vals, 1):
                cell = style_data_cell(ws, dr, ci, alt)
                cell.value = v
                if ci == 4 and v == 'delivered':
                    cell.fill = green_fill
                    cell.font = Font(name="Calibri", size=10, color=XL_WHITE, bold=True)
            dr += 1

    for i, w in enumerate([16, 26, 50, 12, 14, 18], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── Risks sheet ──
    ws = wb.create_sheet("Risks")
    ws.sheet_properties.tabColor = XL_RED
    ws['A1'] = "Risks — source of truth (edit here)"
    ws['A1'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    headers = ["ID", "Title", "Category", "Likelihood", "Impact", "Score", "Status", "Owner", "Mitigation", "Contingency"]
    row = 3
    for ci, h in enumerate(headers, 1):
        ws.cell(row=row, column=ci, value=h)
    style_header_row(ws, row, len(headers))

    def risk_score(l, i):
        scores = {'low': 1, 'medium': 2, 'high': 3}
        s = scores.get(l, 1) * scores.get(i, 1)
        if s >= 6: return 'critical' if s == 9 else 'high'
        if s >= 3: return 'medium'
        return 'low'

    for ri, rk in enumerate(risks):
        r = row + 1 + ri
        alt = ri % 2 == 1
        score = risk_score(rk.get('likelihood', 'medium'), rk.get('impact', 'medium'))
        vals = [
            rk['id'],
            rk.get('title', ''),
            rk.get('category', ''),
            rk.get('likelihood', ''),
            rk.get('impact', ''),
            score,
            rk.get('status', ''),
            rk.get('owner_person', ''),
            rk.get('mitigation', ''),
            rk.get('contingency', ''),
        ]
        for ci, v in enumerate(vals, 1):
            cell = style_data_cell(ws, r, ci, alt)
            cell.value = str(v).strip() if v else ''
            # Color score
            if ci == 6:
                if v == 'critical':
                    cell.fill = red_fill
                    cell.font = Font(name="Calibri", size=10, color=XL_WHITE, bold=True)
                elif v == 'high':
                    cell.fill = amber_fill
                    cell.font = Font(name="Calibri", size=10, color=XL_WHITE, bold=True)

    for i, w in enumerate([18, 50, 14, 12, 12, 10, 10, 24, 60, 60], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── Phases sheet ──
    ws = wb.create_sheet("Phases")
    ws.sheet_properties.tabColor = XL_NAVY
    ws['A1'] = "Phases — source of truth (edit here)"
    ws['A1'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    headers = ["Phase", "Name", "Description", "Start Week", "End Week"]
    row = 3
    for ci, h in enumerate(headers, 1):
        ws.cell(row=row, column=ci, value=h)
    style_header_row(ws, row, len(headers))

    phases_data = []
    for p in sorted((STATE / "phases").iterdir()):
        if p.is_dir():
            pm = load_yaml(p / "manifest.yaml")
            phases_data.append(pm)

    for ri, ph in enumerate(phases_data):
        r = row + 1 + ri
        alt = ri % 2 == 1
        vals = [ph.get('id', ''), ph.get('name', ''), ph.get('description', ''),
                ph.get('start_week', ''), ph.get('end_week', '')]
        for ci, v in enumerate(vals, 1):
            cell = style_data_cell(ws, r, ci, alt)
            cell.value = str(v).strip() if v else ''

    for i, w in enumerate([14, 24, 50, 12, 12], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── Monthly Gantt sheet ──
    ws = wb.create_sheet("Monthly Gantt")
    ws.sheet_properties.tabColor = XL_ACCENT
    ws['A1'] = "Monthly Gantt — derived from Milestones"
    ws['A1'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)
    ws['A2'] = "Collapsed view for sharing/printing. Cells are computed from milestone dates."
    ws['A2'].font = Font(name="Calibri", italic=True, size=10)

    months = ["Apr '26", "May '26", "Jun '26", "Jul '26", "Aug '26", "Sep '26",
              "Oct '26", "Nov '26", "Dec '26", "Jan '27", "Feb '27", "Mar '27"]
    month_starts = [
        datetime.date(2026,4,1), datetime.date(2026,5,1), datetime.date(2026,6,1),
        datetime.date(2026,7,1), datetime.date(2026,8,1), datetime.date(2026,9,1),
        datetime.date(2026,10,1), datetime.date(2026,11,1), datetime.date(2026,12,1),
        datetime.date(2027,1,1), datetime.date(2027,2,1), datetime.date(2027,3,1),
    ]
    month_ends = month_starts[1:] + [datetime.date(2027,4,1)]

    headers = ["ID", "Title", "Status"] + months
    row = 4
    for ci, h in enumerate(headers, 1):
        ws.cell(row=row, column=ci, value=h)
    style_header_row(ws, row, len(headers))

    status_fills = {
        'complete': green_fill,
        'in_progress': amber_fill,
        'planned': PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid"),
    }

    for ri, m in enumerate(milestones):
        r = row + 1 + ri
        alt = ri % 2 == 1

        cell = style_data_cell(ws, r, 1, alt)
        cell.value = m['id']
        cell = style_data_cell(ws, r, 2, alt)
        cell.value = m.get('title', '')
        cell = style_data_cell(ws, r, 3, alt)
        cell.value = m.get('status', '')

        ms = datetime.date.fromisoformat(str(m.get('planned_start', '2026-04-24')))
        me = datetime.date.fromisoformat(str(m.get('planned_end', '2026-04-27')))

        for mi in range(12):
            ci = 4 + mi
            cell = style_data_cell(ws, r, ci, alt)
            # Check overlap
            if ms < month_ends[mi] and me >= month_starts[mi]:
                fill = status_fills.get(m.get('status', 'planned'), status_fills['planned'])
                cell.fill = fill
                cell.value = m.get('status', '')[0].upper() if m.get('status') else ''
                cell.font = Font(name="Calibri", size=9, color=XL_WHITE, bold=True)
                cell.alignment = Alignment(horizontal='center', vertical='center')

    ws.column_dimensions['A'].width = 26
    ws.column_dimensions['B'].width = 44
    ws.column_dimensions['C'].width = 12
    for i in range(4, 16):
        ws.column_dimensions[get_column_letter(i)].width = 10

    # ── Dependencies sheet ──
    ws = wb.create_sheet("Dependencies")
    ws.sheet_properties.tabColor = XL_NAVY
    ws['A1'] = "Dependencies — milestone dependency map"
    ws['A1'].font = Font(name="Calibri", bold=True, size=12, color=XL_NAVY)

    headers = ["Milestone", "Status", "% Complete", "Depends on", "Provides to"]
    row = 3
    for ci, h in enumerate(headers, 1):
        ws.cell(row=row, column=ci, value=h)
    style_header_row(ws, row, len(headers))

    # Simple dependency mapping for project-state
    deps = {
        'M01-v2-release': ('—', 'M02, M09'),
        'M02-starter-pack-hardening': ('M01', 'M05, M06, M10'),
        'M03-matrix-editor': ('M01', 'M10'),
        'M04-pack-registry': ('M01', 'M05, M06, M10'),
        'M05-compliance-packs-wave1': ('M02, M04', 'M10'),
        'M06-compliance-packs-wave2': ('M02, M04', 'M10'),
        'M07-multi-project-dashboard': ('M01', 'M10'),
        'M08-orchestrator-scheduling': ('M01', 'M10'),
        'M09-substrate-validation-suite': ('M01', 'M10'),
        'M10-v21-release': ('M02, M03, M04, M08, M09', '—'),
    }

    for ri, m in enumerate(milestones):
        r = row + 1 + ri
        alt = ri % 2 == 1
        dep_on, provides = deps.get(m['id'], ('—', '—'))
        vals = [m['id'], m.get('status', ''), m.get('percent_complete', 0), dep_on, provides]
        for ci, v in enumerate(vals, 1):
            cell = style_data_cell(ws, r, ci, alt)
            cell.value = v

    for i, w in enumerate([28, 12, 14, 30, 30], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ── Legend sheet ──
    ws = wb.create_sheet("Legend")
    ws.sheet_properties.tabColor = XL_GRAY
    ws['A1'] = "Legend"
    ws['A1'].font = Font(name="Calibri", bold=True, size=14, color=XL_NAVY)

    ws['A3'] = "Status values"
    ws['A3'].font = Font(name="Calibri", bold=True, size=11)
    legend = [
        ("planned", "Not yet started", XL_GRAY),
        ("in_progress", "Work underway", XL_AMBER),
        ("complete", "Done", XL_GREEN),
        ("at_risk", "Behind schedule or blocked", XL_RED),
        ("delivered", "Deliverable accepted", XL_GREEN),
    ]
    for i, (status, desc, color) in enumerate(legend):
        r = 4 + i
        ws.cell(row=r, column=1, value=status).font = Font(name="Calibri", size=10, bold=True)
        ws.cell(row=r, column=2, value=desc).font = Font(name="Calibri", size=10)
        ws.cell(row=r, column=3).fill = PatternFill(start_color=color, end_color=color, fill_type="solid")

    ws['A10'] = "Risk scores"
    ws['A10'].font = Font(name="Calibri", bold=True, size=11)
    risk_legend = [
        ("critical", "high × high", XL_RED),
        ("high", "high × medium or medium × high", XL_AMBER),
        ("medium", "any other combination with medium", "F1C40F"),
        ("low", "low × low or low × medium", XL_GREEN),
    ]
    for i, (score, desc, color) in enumerate(risk_legend):
        r = 11 + i
        ws.cell(row=r, column=1, value=score).font = Font(name="Calibri", size=10, bold=True)
        ws.cell(row=r, column=2, value=desc).font = Font(name="Calibri", size=10)
        ws.cell(row=r, column=3).fill = PatternFill(start_color=color, end_color=color, fill_type="solid")

    ws.column_dimensions['A'].width = 16
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['C'].width = 8

    path = OUT / "01-Project-State-Tracker-v1.0.xlsx"
    wb.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# 2. BASELINE REPORTS INDEX (docx)
# ═══════════════════════════════════════════════════════════════

def build_index():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc_heading(doc, "project-state — Baseline Reports Bundle", level=1)

    p = doc.add_paragraph()
    p.add_run("Project: ").bold = True
    p.add_run("project-state v2.0 — Generic operational substrate for multi-stakeholder projects\n")
    p.add_run("Lead: ").bold = True
    p.add_run("David Olsson (Worksona)\n")
    p.add_run("Date: ").bold = True
    p.add_run(f"{TODAY}\n")
    p.add_run("Phase: ").bold = True
    p.add_run("04-release (agile-default preset)")

    doc_heading(doc, "What this bundle is", level=2)
    doc_para(doc, (
        "Seven documents that together form the baseline view of project-state at v2.0 release — "
        "the shared reference for development planning, risk management, and roadmap tracking. "
        "Every document is generated from the .project-state/ substrate and can be regenerated at any time."
    ))

    doc_heading(doc, "Contents", level=2)

    contents = [
        ("00", "Baseline Reports Index (this document)", "docx", "Cover + table of contents + usage guide"),
        ("01", "Project-State-Tracker-v1.0", "xlsx", "Consolidated tracker: milestones, deliverables, risks, Gantt, dependencies"),
        ("02", "Project-Plan-and-Timeline", "docx", "Project plan, phases, timeline, and milestone overview"),
        ("03", "Risk-Register", "docx", "Structured risk register with scoring matrix and individual entries"),
        ("04", "Milestone-Detailed-Specs", "docx", "Detailed milestone specifications with deliverables and acceptance criteria"),
        ("05", "Architecture-Overview", "docx", "Three-layer architecture, skill inventory, substrate schema"),
        ("06", "Roadmap-and-KPIs", "docx", "Roadmap timeline, KPIs, success metrics, and pack maturity targets"),
    ]

    add_styled_table(doc, ["#", "Document", "Format", "Purpose"], contents)

    doc_heading(doc, "Version history", level=2)
    add_styled_table(doc, ["Version", "Date", "Change summary"], [
        ("0.1", TODAY, "Initial baseline: 7 documents generated from .project-state/ substrate at v2.0 release"),
    ])

    path = OUT / "00-Baseline-Reports-Index.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# 3. PROJECT PLAN AND TIMELINE (docx)
# ═══════════════════════════════════════════════════════════════

def build_project_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc_heading(doc, "project-state — Project Plan and Timeline", level=1)

    p = doc.add_paragraph()
    p.add_run(f"Date: {TODAY}  ·  Phase: 04-release  ·  Pack: agile-default")

    doc_heading(doc, "Project overview", level=2)
    doc_para(doc, manifest['project']['one_liner'].strip())

    doc_heading(doc, "Lifecycle phases", level=2)
    doc_para(doc, (
        "project-state uses the agile-default phase preset with four phases. "
        "The project is currently in phase 04-release following the v2.0 launch."
    ))

    phases_data = []
    for p in sorted((STATE / "phases").iterdir()):
        if p.is_dir():
            pm = load_yaml(p / "manifest.yaml")
            phases_data.append((pm.get('id', ''), pm.get('name', ''), pm.get('description', '')))

    add_styled_table(doc, ["Phase ID", "Name", "Description"], phases_data)

    doc_heading(doc, "Milestone summary", level=2)
    doc_para(doc, (
        f"The project has {len(milestones)} milestones spanning from April 2026 through December 2026. "
        f"M01 (v2.0 release) is complete. Nine milestones remain planned for v2.1 and beyond."
    ))

    rows = []
    for m in milestones:
        rows.append((
            m['id'],
            m.get('title', ''),
            m.get('status', ''),
            f"{m.get('percent_complete', 0)}%",
            str(m.get('planned_start', '')),
            str(m.get('planned_end', '')),
        ))

    add_styled_table(doc, ["ID", "Title", "Status", "% Complete", "Start", "End"], rows)

    doc_heading(doc, "Timeline", level=2)
    doc_para(doc, (
        "Apr 2026: M01 v2.0 Release (COMPLETE)\n"
        "May–Jun 2026: M02 Starter Pack Hardening, M09 Validation Suite\n"
        "Jun–Aug 2026: M03 Matrix Editor, M08 Orchestrator Scheduling\n"
        "Jul–Aug 2026: M04 Pack Registry\n"
        "Aug–Sep 2026: M10 v2.1 Release\n"
        "Aug–Oct 2026: M05 Compliance Packs Wave 1\n"
        "Sep–Nov 2026: M07 Multi-Project Dashboard\n"
        "Oct–Dec 2026: M06 Compliance Packs Wave 2"
    ))

    doc_heading(doc, "Critical path", level=2)
    doc_para(doc, (
        "M01 → M02/M09 → M03/M04/M08 → M10 (v2.1 release)\n"
        "M02/M04 → M05/M06 (compliance packs depend on hardened starters + registry)"
    ))

    doc_heading(doc, "Surfaces and reporting", level=2)
    surfaces = [
        ("Blog (scsiwyg)", "Enabled", "Weekly status, sprint retros"),
        ("Website (Vercel)", "Enabled", "Reference docs, auto-deploy"),
        ("Slack", "Disabled", "—"),
        ("Gmail", "Disabled", "—"),
        ("Calendar", "Disabled", "—"),
    ]
    add_styled_table(doc, ["Surface", "Status", "Usage"], surfaces)

    path = OUT / "02-Project-Plan-and-Timeline.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# 4. RISK REGISTER (docx)
# ═══════════════════════════════════════════════════════════════

def build_risk_register():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc_heading(doc, "project-state — Risk Register (Baseline)", level=1)

    p = doc.add_paragraph()
    p.add_run(f"Date: {TODAY}  ·  Status: Initial register  ·  Owner: David Olsson (Worksona)")

    doc_para(doc, (
        f"This register contains {len(risks)} identified risks. All are currently open. "
        "Updated on each milestone completion and whenever a risk materializes or is mitigated."
    ))

    doc_heading(doc, "Scoring matrix", level=2)
    scoring = [
        ("low", "low", "low"),
        ("low", "medium", "low"),
        ("low", "high", "medium"),
        ("medium", "low", "low"),
        ("medium", "medium", "medium"),
        ("medium", "high", "high"),
        ("high", "low", "medium"),
        ("high", "medium", "high"),
        ("high", "high", "critical"),
    ]
    add_styled_table(doc, ["Likelihood", "Impact", "Score"], scoring)

    doc_heading(doc, "Register at a glance", level=2)

    def score_label(l, i):
        scores = {'low': 1, 'medium': 2, 'high': 3}
        s = scores.get(l, 1) * scores.get(i, 1)
        if s >= 6: return 'critical' if s == 9 else 'high'
        if s >= 3: return 'medium'
        return 'low'

    summary_rows = []
    for r in risks:
        summary_rows.append((
            r['id'],
            r.get('title', ''),
            r.get('likelihood', ''),
            r.get('impact', ''),
            score_label(r.get('likelihood', ''), r.get('impact', '')),
            r.get('owner_person', ''),
        ))

    add_styled_table(doc, ["ID", "Title", "Likelihood", "Impact", "Score", "Owner"], summary_rows)

    doc_heading(doc, "Top 3 for weekly reporting", level=2)
    doc_para(doc, "R-01 — Single maintainer (critical; bus factor of one).")
    doc_para(doc, "R-03 — No automated tests or CI (critical; silent regression risk).")
    doc_para(doc, "R-07 — Timeline overload (critical; 9 milestones for one person).")

    doc_heading(doc, "Individual risk entries", level=1)

    for r in risks:
        doc_heading(doc, f"{r['id']} — {r.get('title', '')}", level=2)

        score = score_label(r.get('likelihood', ''), r.get('impact', ''))
        p = doc.add_paragraph()
        p.add_run(f"Category: ").bold = True
        p.add_run(f"{r.get('category', 'general')}  ·  ")
        p.add_run(f"Likelihood: ").bold = True
        p.add_run(f"{r.get('likelihood', '')}  ·  ")
        p.add_run(f"Impact: ").bold = True
        p.add_run(f"{r.get('impact', '')}  ·  ")
        p.add_run(f"Score: ").bold = True
        p.add_run(f"{score}")

        p2 = doc.add_paragraph()
        p2.add_run(f"Owner: ").bold = True
        p2.add_run(str(r.get('owner_person', '')))
        if r.get('related_milestones'):
            p2.add_run(f"\nRelated milestones: ").bold = True
            p2.add_run(str(r.get('related_milestones', '')))

        doc_heading(doc, "Description", level=3)
        doc_para(doc, str(r.get('description', '')).strip())

        doc_heading(doc, "Mitigation", level=3)
        doc_para(doc, str(r.get('mitigation', '')).strip())

        if r.get('contingency'):
            doc_heading(doc, "Contingency", level=3)
            doc_para(doc, str(r.get('contingency', '')).strip())

    path = OUT / "03-Risk-Register.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# 5. MILESTONE DETAILED SPECS (docx)
# ═══════════════════════════════════════════════════════════════

def build_milestone_specs():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc_heading(doc, "project-state — Milestone Detailed Specifications", level=1)

    p = doc.add_paragraph()
    p.add_run(f"Date: {TODAY}  ·  Milestones: {len(milestones)}  ·  Deliverables: "
              f"{sum(len(m.get('deliverables', [])) for m in milestones)}")

    for m in milestones:
        doc_heading(doc, f"{m['id']} — {m.get('title', '')}", level=2)

        p = doc.add_paragraph()
        p.add_run("Status: ").bold = True
        p.add_run(f"{m.get('status', '')}  ·  ")
        p.add_run("% Complete: ").bold = True
        p.add_run(f"{m.get('percent_complete', 0)}%  ·  ")
        p.add_run("Owner: ").bold = True
        p.add_run(str(m.get('owner_person', '')))

        p2 = doc.add_paragraph()
        p2.add_run("Planned: ").bold = True
        p2.add_run(f"{m.get('planned_start', '')} → {m.get('planned_end', '')}")
        if m.get('actual_start'):
            p2.add_run(f"  ·  Actual: {m.get('actual_start', '')} → {m.get('actual_end', '')}")

        doc_heading(doc, "Description", level=3)
        doc_para(doc, str(m.get('description', '')).strip())

        if m.get('technical_progress'):
            doc_heading(doc, "Technical progress", level=3)
            doc_para(doc, str(m.get('technical_progress', '')).strip())

        doc_heading(doc, "Deliverables", level=3)
        delivs = [(d['id'], d['title'], d['status']) for d in m.get('deliverables', [])]
        if delivs:
            add_styled_table(doc, ["ID", "Title", "Status"], delivs)
        else:
            doc_para(doc, "No deliverables defined.")

    path = OUT / "04-Milestone-Detailed-Specs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# 6. ARCHITECTURE OVERVIEW (docx)
# ═══════════════════════════════════════════════════════════════

def build_architecture():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc_heading(doc, "project-state — Architecture Overview", level=1)

    p = doc.add_paragraph()
    p.add_run(f"Date: {TODAY}  ·  Version: v2.0")

    doc_heading(doc, "Three-layer architecture", level=2)

    layers = [
        ("Substrate", ".project-state/", "File-per-entity storage with append-only activity logs, advisory lockfiles, and frontmatter timestamps for concurrency. Plain YAML, JSON, NDJSON, and markdown — no database or API."),
        ("Skills", "skills/project-*/SKILL.md", "18 markdown-defined skills grouped into P0 (foundations), P1 (core ops), P2 (surfaces/automation), P3 (polish). Skills are stateless verbs that read/write the substrate via project-state."),
        ("Surfaces", "Slack, Gmail, Calendar, Blog, Website", "External delivery channels. All artifacts require human review before sending. Gmail is always draft. Calendar events are proposed holds."),
    ]
    add_styled_table(doc, ["Layer", "Location", "Description"], layers)

    doc_heading(doc, "Skill inventory", level=2)

    skills = [
        ("P0 Foundation", "project-state", "Memory layer — only skill that reads/writes .project-state/ directly"),
        ("P0 Foundation", "project-scaffolder", "One-shot initializer for new .project-state/"),
        ("P1 Core", "project-phase-gate", "Lifecycle phase transitions with gate enforcement"),
        ("P1 Core", "project-document-curator", "Classify, index, and promote project documents"),
        ("P1 Core", "project-milestone-manager", "CRUD milestones, percent_complete, technical_progress"),
        ("P1 Core", "project-status-reporter", "Generate status reports (weekly, SC pack, claim draft, dashboard)"),
        ("P2 Surfaces", "project-orchestrator", "Calendar-aware conductor; reads reporting matrix, dispatches generators"),
        ("P2 Surfaces", "project-notifier", "Route artifacts to Slack/Gmail/Calendar"),
        ("P2 Surfaces", "project-review-meeting", "Generic recurring review-meeting lifecycle"),
        ("P2 Surfaces", "project-funder-reporting", "Generic stakeholder-bound recurring reports"),
        ("P2 Surfaces", "project-change-register", "Material vs. non-material change classification"),
        ("P2 Surfaces", "project-blog-publisher", "scsiwyg bridge with publication-review respect"),
        ("P2 Surfaces", "project-website-publisher", "Static project website with stable URLs"),
        ("P3 Polish", "project-onboarder", "Personalized onboarding briefs for new teammates"),
        ("P3 Polish", "project-ip-tracker", "IP disclosures with configurable recipient"),
        ("P3 Polish", "project-external-comms", "External-communication review pipeline"),
        ("P3 Polish", "project-lessons", "Continuous lessons-learned capture"),
        ("P3 Polish", "project-archive", "Project closeout and archival"),
    ]
    add_styled_table(doc, ["Tier", "Skill", "Purpose"], skills)

    doc_heading(doc, "Pack system", level=2)
    doc_para(doc, (
        "Packs are pluggable compliance/reporting configurations. Each pack contains a manifest.yaml, "
        "profile YAMLs, optional templates, and reporting-matrix-defaults.yaml. Six skills are pack-profile-driven."
    ))

    packs = [
        ("pic-pcais", "production", "Protein Industries Canada PCAIS program"),
        ("client-services", "starter", "Client services / consulting engagements"),
        ("board-investor", "starter", "Board and investor reporting"),
        ("agile-default", "starter", "Agile / open-source development"),
        ("open-source-community", "starter", "Open-source community governance"),
    ]
    add_styled_table(doc, ["Pack", "Maturity", "Description"], packs)

    doc_heading(doc, "Concurrency model", level=2)
    doc_para(doc, (
        "File-per-entity prevents write conflicts on different entities. "
        "Monolithic files (manifest.yaml, state.json) use advisory lockfiles with 300s TTL. "
        "Frontmatter last_modified timestamps for optimistic concurrency. "
        "Logs are append-only; corrections are new entries, not rewrites."
    ))

    doc_heading(doc, "Reporting matrix", level=2)
    doc_para(doc, (
        "The central v2.0 concept: .project-state/reporting-matrix.yaml encodes "
        "\"for each stakeholder group, what report at what cadence in what format on which surface, "
        "produced by which skill.\" The orchestrator reads this matrix on each tick and dispatches generators."
    ))

    path = OUT / "05-Architecture-Overview.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# 7. ROADMAP AND KPIs (docx)
# ═══════════════════════════════════════════════════════════════

def build_roadmap_kpis():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc_heading(doc, "project-state — Roadmap and KPIs", level=1)

    p = doc.add_paragraph()
    p.add_run(f"Date: {TODAY}  ·  Horizon: Apr 2026 — Dec 2026")

    doc_heading(doc, "Roadmap phases", level=2)

    phases = [
        ("v2.0 (Done)", "Apr 2026", "18 skills, 5 packs, migration script, full docs, project website"),
        ("v2.0.x Hardening", "May–Jun 2026", "Starter pack hardening (M02), validation suite (M09)"),
        ("v2.1 Features", "Jun–Sep 2026", "Matrix editor (M03), pack registry (M04), orchestrator scheduling (M08)"),
        ("v2.1 Release", "Aug–Sep 2026", "Cut v2.1 with self-driving orchestrator + matrix editor"),
        ("Compliance Expansion", "Aug–Dec 2026", "Wave 1 academic/federal (M05), Wave 2 industry/regulatory (M06)"),
        ("Portfolio", "Sep–Nov 2026", "Multi-project dashboard (M07)"),
    ]
    add_styled_table(doc, ["Phase", "Timeline", "Scope"], phases)

    doc_heading(doc, "Key Performance Indicators", level=2)

    kpis = [
        ("Pack maturity", "4 starter packs promoted to production", "0/4", "M02"),
        ("Test coverage", "Schema validators + pre-commit hooks in CI", "0%", "M09"),
        ("Self-driving ops", "Orchestrator fires on cron, no human initiation", "Not started", "M08"),
        ("Non-technical access", "Matrix editor usable by non-technical PLs", "Not started", "M03"),
        ("Pack distribution", "Install packs by name from registry", "Not started", "M04"),
        ("Compliance breadth", "3+ additional compliance packs (beyond pic-pcais)", "0/3", "M05, M06"),
        ("Portfolio support", "Dashboard across 3+ project instances", "Not started", "M07"),
        ("Community", "Public repo + CONTRIBUTING.md + first external pack", "Not started", "—"),
    ]
    add_styled_table(doc, ["KPI", "Target", "Current", "Milestone"], kpis)

    doc_heading(doc, "Success criteria for v2.1", level=2)
    doc_para(doc, (
        "1. All four starter packs at production maturity (validated on real projects).\n"
        "2. Substrate validation suite catches schema drift and cross-reference breaks.\n"
        "3. Orchestrator runs on cron and generates draft reports without human initiation.\n"
        "4. Matrix editor allows non-technical PLs to add/edit reporting entries.\n"
        "5. At least one compliance pack from Wave 1 reaches production maturity."
    ))

    doc_heading(doc, "Pack maturity targets", level=2)

    pack_targets = [
        ("pic-pcais", "production", "production", "Maintain"),
        ("client-services", "starter", "production", "M02"),
        ("board-investor", "starter", "production", "M02"),
        ("agile-default", "starter", "production", "M02"),
        ("open-source-community", "starter", "production", "M02"),
        ("nserc-academic", "—", "production", "M05"),
        ("nih-rppr", "—", "production", "M05"),
        ("nsf-merit", "—", "production", "M05"),
        ("soc2-evidence", "—", "starter+", "M06"),
        ("iso-27001", "—", "starter+", "M06"),
        ("eu-horizon", "—", "starter+", "M06"),
    ]
    add_styled_table(doc, ["Pack", "Current", "Target", "Milestone"], pack_targets)

    path = OUT / "06-Roadmap-and-KPIs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")

# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print(f"Generating baseline reports for project-state v2.0 ({TODAY})...")
    print(f"Output: {OUT}/\n")
    build_tracker()
    build_index()
    build_project_plan()
    build_risk_register()
    build_milestone_specs()
    build_architecture()
    build_roadmap_kpis()
    print(f"\nDone. 7 files generated.")
